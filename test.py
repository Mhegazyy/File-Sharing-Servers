from docx import Document

# Initialize the document
doc = Document()
doc.add_heading('Testing and Results', level=1)

# Introduction
doc.add_paragraph(
    "The file sharing server underwent extensive testing to ensure its functionality, "
    "performance, and security. This section outlines the testing methodologies, key results, "
    "and supporting evidence for each test case."
)

# Functional Testing
doc.add_heading('1. Functional Testing', level=2)
doc.add_paragraph(
    "The following table summarizes the functional test cases, expected outcomes, "
    "actual results, and includes placeholders for relative hyperlinks to screenshots."
)

# Add Functional Testing Table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Test Case Description'
hdr_cells[2].text = 'Expected Outcome'
hdr_cells[3].text = 'Actual Outcome'
hdr_cells[4].text = 'Evidence'

# Test Cases with Relative Paths to Screenshots
functional_tests = [
    ["User Registration", "Register a new user with valid credentials.", "User registered successfully.", "Pass", "screenshots/Screenshot1.png"],
    ["User Registration", "Register with an existing username.", "Duplicate registration rejected.", "Pass", "screenshots/Screenshot2.png"],
    ["User Login", "Login with valid credentials.", "Session token returned.", "Pass", "screenshots/Screenshot3.png"],
    ["User Login", "Login with incorrect credentials.", "Access denied with error message.", "Pass", "screenshots/Screenshot4.png"],
    ["File Upload", "Upload a valid file.", "File encrypted and saved successfully.", "Pass", "screenshots/Screenshot5.png"],
    ["File Upload", "Upload with invalid session token.", "Upload denied with session error.", "Pass", "screenshots/Screenshot6.png"],
    ["File Download", "Download an existing file.", "Encrypted file sent and decrypted.", "Pass", "screenshots/Screenshot7.png"],
    ["File Download", "Download a non-existing file.", "File not found error.", "Pass", "screenshots/Screenshot8.png"],
    ["File Integrity", "Upload a file and compare checksums.", "Checksums match; integrity validated.", "Pass", "screenshots/Screenshot9.png"],
    ["File Integrity", "Modify file during transfer.", "Checksum mismatch; transfer rejected.", "Pass", "screenshots/Screenshot10.png"],
]

# Populate the table
for test in functional_tests:
    row_cells = table.add_row().cells
    row_cells[0].text = test[0]
    row_cells[1].text = test[1]
    row_cells[2].text = test[2]
    row_cells[3].text = test[3]
    # Add a hyperlink for the Evidence column
    paragraph = row_cells[4].paragraphs[0]
    run = paragraph.add_run("View Screenshot")
    run.font.underline = True
    paragraph.hyperlink = test[4]

# Save the document
output_file = "testing_results.docx"
doc.save(output_file)

print(f"Document saved successfully as {output_file}")
